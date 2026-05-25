from pathlib import Path
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from docx import Document
from docx.shared import Pt, RGBColor


_DARK = HexColor("#111111")
_MID = HexColor("#333333")
_MUTED = HexColor("#777777")
_RULE = HexColor("#dddddd")


def _build_styles():
    return {
        "name": ParagraphStyle(
            "Name", fontSize=20, fontName="Helvetica-Bold",
            spaceAfter=6, textColor=_DARK,
        ),
        "contact": ParagraphStyle(
            "Contact", fontSize=9.5, fontName="Helvetica",
            spaceAfter=0, textColor=_MUTED, leading=14,
        ),
        "section": ParagraphStyle(
            "Section", fontSize=9.5, fontName="Helvetica-Bold",
            spaceBefore=12, spaceAfter=4, textColor=_DARK,
            letterSpacing=1.2,
        ),
        "job_title": ParagraphStyle(
            "JobTitle", fontSize=10.5, fontName="Helvetica-Bold",
            spaceAfter=1, textColor=_DARK,
        ),
        "job_meta": ParagraphStyle(
            "JobMeta", fontSize=9, fontName="Helvetica",
            spaceAfter=3, textColor=_MUTED,
        ),
        "normal": ParagraphStyle(
            "Normal2", fontSize=9.5, fontName="Helvetica",
            spaceAfter=1, leading=13, textColor=_MID,
        ),
        "bullet": ParagraphStyle(
            "Bullet", fontSize=9.5, fontName="Helvetica",
            leftIndent=10, firstLineIndent=0,
            spaceAfter=2, leading=13, textColor=_MID,
        ),
    }


def _rule(space_before=0, space_after=6):
    return HRFlowable(
        width="100%", thickness=0.4, color=_RULE,
        spaceBefore=space_before, spaceAfter=space_after,
    )


def _clean_linkedin(url: str) -> str:
    if not url:
        return ""
    url = url.rstrip("/")
    if "linkedin.com/in/" in url:
        profile_id = url.split("linkedin.com/in/")[-1]
        return f"linkedin.com/in/{profile_id}"
    return url


def _clean_github(url: str) -> str:
    if not url:
        return ""
    url = url.rstrip("/")
    if "github.com/" in url:
        return "github.com/" + url.split("github.com/")[-1]
    return url


def to_pdf(data: dict, output_path: str) -> str:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        output_path, pagesize=letter,
        leftMargin=0.7 * inch, rightMargin=0.7 * inch,
        topMargin=0.6 * inch, bottomMargin=0.6 * inch,
    )
    s = _build_styles()
    story = []

    # ── Header ──────────────────────────────────────────────────────────────
    contact = data.get("contact", {})
    story.append(Paragraph(contact.get("name", "Candidate"), s["name"]))
    story.append(_rule(space_before=4, space_after=5))

    parts = [x for x in [
        contact.get("email"),
        contact.get("phone"),
        _clean_linkedin(contact.get("linkedin", "")),
        _clean_github(contact.get("github", "")),
        contact.get("location"),
    ] if x]
    story.append(Paragraph("  ·  ".join(parts), s["contact"]))
    story.append(Spacer(1, 10))

    def section(title):
        story.append(Paragraph(title.upper(), s["section"]))
        story.append(_rule(space_before=2, space_after=4))

    # ── Summary ──────────────────────────────────────────────────────────────
    if data.get("summary"):
        section("Summary")
        story.append(Paragraph(data["summary"], s["normal"]))
        story.append(Spacer(1, 2))

    # ── Skills ───────────────────────────────────────────────────────────────
    if data.get("skills"):
        section("Skills")
        story.append(Paragraph(", ".join(data["skills"]), s["normal"]))
        story.append(Spacer(1, 2))

    # ── Experience ───────────────────────────────────────────────────────────
    if data.get("experience"):
        section("Experience")
        for exp in data["experience"]:
            story.append(Paragraph(exp.get("title", ""), s["job_title"]))
            meta_parts = [p for p in [exp.get("company"), exp.get("duration")] if p]
            story.append(Paragraph("  ·  ".join(meta_parts), s["job_meta"]))
            for bullet in exp.get("bullets", []):
                story.append(Paragraph(f"–  {bullet}", s["bullet"]))
            story.append(Spacer(1, 6))

    # ── Projects ─────────────────────────────────────────────────────────────
    if data.get("projects"):
        section("Projects")
        for proj in data["projects"]:
            techs = ", ".join(proj.get("technologies", []))
            story.append(Paragraph(proj.get("name", ""), s["job_title"]))
            if techs:
                story.append(Paragraph(techs, s["job_meta"]))
            story.append(Paragraph(proj.get("description", ""), s["normal"]))
            story.append(Spacer(1, 5))

    # ── Education ────────────────────────────────────────────────────────────
    if data.get("education"):
        section("Education")
        for edu in data["education"]:
            story.append(Paragraph(edu.get("degree", ""), s["job_title"]))
            meta_parts = [p for p in [edu.get("institution"), edu.get("year")] if p]
            story.append(Paragraph("  ·  ".join(meta_parts), s["job_meta"]))

    # ── Certifications ────────────────────────────────────────────────────────
    if data.get("certifications"):
        section("Certifications")
        for cert in data["certifications"]:
            story.append(Paragraph(f"–  {cert}", s["bullet"]))

    doc.build(story)
    return output_path


def to_docx(data: dict, output_path: str) -> str:
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    contact = data.get("contact", {})
    doc.add_heading(contact.get("name", "Candidate"), 0)
    parts = [x for x in [
        contact.get("email"), contact.get("phone"),
        _clean_linkedin(contact.get("linkedin", "")),
        _clean_github(contact.get("github", "")),
    ] if x]
    doc.add_paragraph("  ·  ".join(parts))

    if data.get("summary"):
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(", ".join(data["skills"]))

    if data.get("experience"):
        doc.add_heading("Experience", level=1)
        for exp in data["experience"]:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(exp.get("title", ""))
            run.bold = True
            meta = [x for x in [exp.get("company"), exp.get("duration")] if x]
            if meta:
                p.add_run(f"  ·  {'  ·  '.join(meta)}")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(f"– {bullet}")

    if data.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            techs = ", ".join(proj.get("technologies", []))
            if techs:
                p.add_run(f"  ·  {techs}")
            doc.add_paragraph(proj.get("description", ""))

    if data.get("education"):
        doc.add_heading("Education", level=1)
        for edu in data["education"]:
            p = doc.add_paragraph()
            p.add_run(edu.get("degree", "")).bold = True
            meta = [x for x in [edu.get("institution"), edu.get("year")] if x]
            if meta:
                p.add_run(f"  ·  {'  ·  '.join(meta)}")

    if data.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in data["certifications"]:
            doc.add_paragraph(f"– {cert}")

    doc.save(output_path)
    return output_path
